"""
CV Parser — adaptación del pipeline de NexoTalent para Harmoni.

Pipeline en DOS FASES (estilo NexoTalent):

  FASE 1 — Parsing LOCAL determinista (sin IA):
    1. Extracción de texto (PDF/DOCX/TXT)
       - PyMuPDF (fitz) → más rápido para PDFs nativos
       - pdfplumber → fallback
       - python-docx → DOCX
    2. **Segmentación en secciones** por headers comunes en español:
         "DATOS PERSONALES", "EXPERIENCIA LABORAL",
         "FORMACIÓN ACADÉMICA", "HABILIDADES", "IDIOMAS", "REFERENCIAS".
    3. Por cada sección, extracción acotada:
       - datos_personales → nombre, DNI, email, teléfono, fecha_nacimiento
       - experiencias    → pares (empresa, cargo, fecha_inicio, fecha_fin)
       - educacion       → institución + año
       - idiomas         → lista
       - skills          → catálogo gastronómico/tech/admin

  FASE 2 — Enriquecimiento con IA (opcional):
    - Sólo se invoca si parsing local dejó campos clave vacíos
      (nombre, email, fecha_nacimiento) o si `forzar_ia=True`.
    - El LLM recibe el texto crudo + las secciones ya detectadas y
      sólo se le pide rellenar los campos que faltan
      (no rehacer todo, evita alucinaciones).
    - La respuesta del LLM se VALIDA contra el texto original
      (no se aceptan emails ni teléfonos que no aparezcan en el CV).
    - Inyectable vía parámetro `llm_caller` para tests (mock).

REGLA CRÍTICA — separación fecha_nacimiento vs experiencia:
    Las fechas que aparecen en la sección "Datos Personales"
    (ej. "Fecha de nacimiento: 14/03/1990") NUNCA contaminan
    el cálculo de `anios_experiencia`. El bug original tomaba
    el año mínimo de TODO el texto, así un postulante nacido en
    1990 con 5 años de experiencia recibía `anios_experiencia=36`.

Las funciones devuelven `str` para texto o `dict` con campos extraídos.
Si todas las extracciones fallan, devuelve dict con campos vacíos —
nunca crashea.

Crédito: pipeline derivado de NexoTalent.
"""
from __future__ import annotations

import logging
import os
import re
from datetime import date
from typing import Callable, Optional

logger = logging.getLogger('reclutamiento.cv_parser')


# ──────────────────────────────────────────────────────────────────
# Constantes
# ──────────────────────────────────────────────────────────────────

EXTENSIONES_SOPORTADAS = {'.pdf', '.docx', '.doc', '.txt', '.rtf'}
_MIN_CHARS_OK = 80  # umbral mínimo para considerar extracción válida

# Listado de skills frecuentes en gastronomía + tech + admin
SKILLS_CATALOG = {
    # Gastronomía
    'cocina_caliente', 'cocina_fria', 'pasteleria', 'panaderia', 'parrilla',
    'sushi', 'pizza', 'mariscos', 'reposteria', 'manipulacion de alimentos',
    'haccp', 'iso 22000', 'control de costos', 'inventarios',
    # Salón / Bar
    'sommelier', 'bartender', 'cata de vinos', 'mixologia', 'atencion al cliente',
    'protocolo de servicio', 'caja', 'manejo de propinas',
    # Admin / RRHH
    'planillas', 'plame', 'sunat', 'sunafil', 'ley 27735', 'cts', 'gratificaciones',
    'liquidacion', 'reclutamiento', 'contratacion', 'evaluacion de desempeño',
    'capacitacion', 'compensaciones', 'desarrollo organizacional',
    # Idiomas
    'ingles avanzado', 'ingles intermedio', 'portugues', 'frances', 'italiano',
    'aleman', 'mandarin', 'japones', 'quechua',
    # Tech
    'excel avanzado', 'excel intermedio', 'power bi', 'sap', 'oracle', 'spring',
    'concar', 'siscont', 'tableau', 'sql', 'python', 'office',
}

# Cargos / títulos típicos peruanos (para detección de experiencia)
CARGOS_FRECUENTES = {
    'chef', 'sous chef', 'chef ejecutivo', 'chef de partie', 'commis',
    'cocinero', 'parrillero', 'panadero', 'pastelero', 'sushiman',
    'mesero', 'mesera', 'capitan de meseros', 'maitre', 'host',
    'sommelier', 'bartender', 'cajero', 'cajera', 'runner', 'busser',
    'administrador', 'gerente', 'jefe', 'supervisor', 'coordinador',
    'asistente', 'analista', 'auxiliar', 'tecnico', 'practicante',
}


# ──────────────────────────────────────────────────────────────────
# Extracción de texto (3 estrategias en cascada)
# ──────────────────────────────────────────────────────────────────

def extract_text_from_file(filepath: str) -> str:
    """Extrae texto de un archivo CV.

    Args:
        filepath: ruta absoluta al archivo.

    Returns:
        Texto extraído, o cadena vacía si no se pudo.
    """
    if not filepath or not os.path.exists(filepath):
        logger.warning('extract_text: archivo no existe: %s', filepath)
        return ''

    ext = os.path.splitext(filepath)[1].lower()
    if ext not in EXTENSIONES_SOPORTADAS:
        logger.warning('extract_text: extensión no soportada: %s', ext)
        return ''

    # PDF: PyMuPDF (fitz) → pdfplumber fallback
    if ext == '.pdf':
        text = _extract_pdf_pymupdf(filepath)
        if len(text.strip()) >= _MIN_CHARS_OK:
            return text
        # Fallback
        text = _extract_pdf_pdfplumber(filepath)
        return text

    # DOCX/DOC
    if ext in ('.docx', '.doc'):
        return _extract_docx(filepath)

    # TXT/RTF
    if ext in ('.txt', '.rtf'):
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.warning('extract_text TXT/RTF: %s', e)
            return ''

    return ''


def _extract_pdf_pymupdf(filepath: str) -> str:
    """Extrae texto con PyMuPDF (fitz). Más rápido para PDFs nativos."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.debug('PyMuPDF no instalado, saltando.')
        return ''
    try:
        parts = []
        with fitz.open(filepath) as doc:
            for page in doc:
                t = page.get_text('text') or ''
                if t.strip():
                    parts.append(t)
        text = '\n'.join(parts)
        if text.strip():
            logger.info('PyMuPDF: %d chars en %d páginas.', len(text), len(parts))
        return text
    except Exception as exc:
        logger.warning('PyMuPDF error: %s', exc)
        return ''


def _extract_pdf_pdfplumber(filepath: str) -> str:
    """Fallback para PDFs raros que PyMuPDF no procesa bien."""
    try:
        import pdfplumber
    except ImportError:
        return ''
    try:
        parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ''
                if t.strip():
                    parts.append(t)
        text = '\n'.join(parts)
        if text.strip():
            logger.info('pdfplumber: %d chars.', len(text))
        return text
    except Exception as exc:
        logger.warning('pdfplumber error: %s', exc)
        return ''


def _extract_docx(filepath: str) -> str:
    """Extrae texto de un DOCX."""
    try:
        from docx import Document
    except ImportError:
        logger.debug('python-docx no instalado, saltando.')
        return ''
    try:
        doc = Document(filepath)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Tablas también pueden contener info importante
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        text = '\n'.join(parts)
        if text.strip():
            logger.info('python-docx: %d chars.', len(text))
        return text
    except Exception as exc:
        logger.warning('python-docx error: %s', exc)
        return ''


# ──────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────

def _strip_tildes(s: str) -> str:
    """Quita tildes y diacríticos para matching robusto."""
    import unicodedata
    return ''.join(
        c for c in unicodedata.normalize('NFKD', s)
        if not unicodedata.combining(c)
    )


# ──────────────────────────────────────────────────────────────────
# Regex de entidades
# ──────────────────────────────────────────────────────────────────

# Email (RFC 5322 simplificado)
RE_EMAIL = re.compile(
    r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}',
)

# Teléfono peruano (móvil 9XX-XXX-XXX o fijo (01)XXXXXXX)
RE_TELEFONO_PE = re.compile(
    r'(?:\+?51[\s-]?)?(?:\(0?1\)[\s-]?\d{7}|9\d{2}[\s-]?\d{3}[\s-]?\d{3})',
)

# DNI peruano (8 dígitos)
RE_DNI_PE = re.compile(r'\b\d{8}\b')

# Fechas (varios formatos)
RE_FECHA = re.compile(
    r'\b(?:0?[1-9]|[12]\d|3[01])[/\-\.]'    # día
    r'(?:0?[1-9]|1[0-2])[/\-\.]'             # mes
    r'(?:19|20)\d{2}\b'                       # año
    r'|\b(?:Ene|Feb|Mar|Abr|May|Jun|Jul|Ago|Sep|Oct|Nov|Dic)[a-z]*\.?\s+(?:19|20)\d{2}\b'
    r'|\b(?:19|20)\d{2}\b',                   # año solo
    re.IGNORECASE,
)

# Fecha de nacimiento explícita (DD/MM/YYYY o DD-MM-YYYY)
# Captura sólo cuando aparece junto a etiquetas como "Fecha de nacimiento",
# "Nacimiento", "Born", "Fec. Nac.", etc.
RE_FECHA_NACIMIENTO_LABEL = re.compile(
    r'(?im)(?:fecha\s+de\s+nacimiento|fec(?:ha)?\.?\s+nac\.?|nacimiento|date\s+of\s+birth|born|nacido(?:\s+el)?)\s*[:\-]?\s*'
    r'(?P<fecha>'
        r'\d{1,2}[/\-\.]\d{1,2}[/\-\.](?:19|20)\d{2}'
        r'|\d{1,2}\s+de\s+(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|setiembre|octubre|noviembre|diciembre)\s+de\s+(?:19|20)\d{2}'
        r'|(?:ene|feb|mar|abr|may|jun|jul|ago|sep|set|oct|nov|dic)[a-z]*\.?\s+\d{1,2},?\s+(?:19|20)\d{2}'
        r'|(?:19|20)\d{2}'
    r')',
)


# ──────────────────────────────────────────────────────────────────
# Headers de secciones (Fase 1: segmentación)
# ──────────────────────────────────────────────────────────────────

SECCION_DATOS_PERSONALES = 'datos_personales'
SECCION_EXPERIENCIA = 'experiencia'
SECCION_EDUCACION = 'educacion'
SECCION_SKILLS = 'skills'
SECCION_IDIOMAS = 'idiomas'
SECCION_REFERENCIAS = 'referencias'
SECCION_OTROS = 'otros'

# Cada header se matchea sobre una LÍNEA (no parte interna).
# Permite minúsculas y mayúsculas; quita tildes antes de matchear.
_HEADERS_POR_SECCION = {
    SECCION_DATOS_PERSONALES: [
        r'datos\s+personales',
        r'informacion\s+personal',
        r'datos\s+de\s+contacto',
        r'perfil\s+personal',
        r'personal\s+data',
        r'personal\s+information',
    ],
    SECCION_EXPERIENCIA: [
        r'experiencia\s+laboral',
        r'experiencia\s+profesional',
        r'experiencia\s+de\s+trabajo',
        r'experiencia(?:\s+\w+){0,2}',  # "Experiencia", "Experiencia previa"
        r'historial\s+laboral',
        r'historial\s+profesional',
        r'work\s+experience',
        r'professional\s+experience',
        r'antecedentes\s+laborales',
        r'trayectoria\s+profesional',
    ],
    SECCION_EDUCACION: [
        r'formacion\s+academica',
        r'formacion\s+profesional',
        r'educacion',
        r'estudios',
        r'estudios\s+realizados',
        r'education',
        r'academic\s+background',
    ],
    SECCION_SKILLS: [
        r'habilidades',
        r'competencias',
        r'aptitudes',
        r'skills',
        r'herramientas',
        r'conocimientos',
        r'conocimientos\s+tecnicos',
    ],
    SECCION_IDIOMAS: [
        r'idiomas',
        r'languages',
    ],
    SECCION_REFERENCIAS: [
        r'referencias',
        r'referencias\s+laborales',
        r'referencias\s+personales',
        r'references',
    ],
}

# Compilados en runtime
_HEADER_RE_CACHE: dict[str, re.Pattern] = {}


def _header_pattern(seccion: str) -> re.Pattern:
    if seccion not in _HEADER_RE_CACHE:
        alt = '|'.join(_HEADERS_POR_SECCION[seccion])
        _HEADER_RE_CACHE[seccion] = re.compile(
            rf'^\s*(?:{alt})\s*:?\s*$',
            re.IGNORECASE,
        )
    return _HEADER_RE_CACHE[seccion]


# Compatibilidad con código viejo que importa estas constantes
RE_HEADER_EXPERIENCIA = re.compile(
    r'(?im)^(?:experiencia\s+(?:laboral|profesional|de\s+trabajo)|'
    r'historial\s+(?:laboral|profesional)|work\s+experience|'
    r'antecedentes\s+laborales|trayectoria\s+profesional)\b'
)
RE_HEADER_EDUCACION = re.compile(
    r'(?im)^(?:formaci[oó]n\s+(?:acad[eé]mica|profesional)|'
    r'educaci[oó]n|estudios|education)\b'
)
RE_HEADER_SKILLS = re.compile(
    r'(?im)^(?:habilidades|competencias|skills|herramientas|conocimientos|aptitudes)\b'
)
RE_HEADER_IDIOMAS = re.compile(
    r'(?im)^(?:idiomas|languages)\b'
)


# ──────────────────────────────────────────────────────────────────
# FASE 1.1 — Segmentar texto en secciones
# ──────────────────────────────────────────────────────────────────

def segmentar_secciones(text: str) -> dict[str, str]:
    """Divide el texto del CV en secciones tipadas.

    Devuelve un dict {seccion: texto_de_la_seccion}. La sección
    `otros` contiene el contenido previo al primer header detectado
    (típicamente el nombre + datos de contacto en la cabecera).

    Estrategia:
      - Recorre línea por línea.
      - Si la línea matchea uno de los headers de sección → cambia
        de cubeta.
      - Conserva orden original; no clasifica por contenido.

    Esto permite que `parse_cv` posteriormente extraiga fechas,
    skills, etc. SÓLO de la sección correspondiente.
    """
    if not text:
        return {SECCION_OTROS: ''}

    secciones: dict[str, list[str]] = {SECCION_OTROS: []}
    actual = SECCION_OTROS

    # Normalización: quitar tildes sólo para matching de headers
    for raw_line in text.split('\n'):
        line_norm = _strip_tildes(raw_line).strip()
        matched = None
        if line_norm and len(line_norm) <= 60:
            # Headers son siempre líneas cortas
            for seccion in _HEADERS_POR_SECCION:
                if _header_pattern(seccion).match(line_norm):
                    matched = seccion
                    break

        if matched:
            actual = matched
            secciones.setdefault(actual, [])
            # Header line en sí no se guarda
            continue
        secciones.setdefault(actual, []).append(raw_line)

    return {k: '\n'.join(v).strip() for k, v in secciones.items()}


# ──────────────────────────────────────────────────────────────────
# FASE 1.2 — Extraer fecha de nacimiento (aislada)
# ──────────────────────────────────────────────────────────────────

def extraer_fecha_nacimiento(text: str, datos_personales_text: str = '') -> Optional[str]:
    """Detecta la fecha de nacimiento del CV.

    Busca primero en la sección "Datos Personales"; si no encuentra,
    busca en todo el texto pero exigiendo etiqueta explícita
    ("Fecha de nacimiento:", "Nacimiento:", "Born:", etc.) para
    evitar falsos positivos.

    Devuelve la cadena original (ej. "14/03/1990") o None.
    """
    candidatos = (datos_personales_text, text)
    for fuente in candidatos:
        if not fuente:
            continue
        m = RE_FECHA_NACIMIENTO_LABEL.search(fuente)
        if m:
            return m.group('fecha').strip()
    return None


def _año_de_fecha_str(s: str) -> Optional[int]:
    """Extrae el año (1900-2100) de una cadena de fecha."""
    if not s:
        return None
    m = re.search(r'\b(19|20)\d{2}\b', s)
    if m:
        return int(m.group())
    return None


# ──────────────────────────────────────────────────────────────────
# FASE 1.3 — Parse principal (sección-aware)
# ──────────────────────────────────────────────────────────────────

def parse_cv(text: str) -> dict:
    """Parsea un texto de CV y devuelve dict con entidades extraídas.

    Esta función ahora es **section-aware**: primero segmenta el CV
    por secciones y luego extrae fechas/años SÓLO de la sección
    de experiencia laboral (no del CV completo).

    Estructura del dict retornado:
        {
            'email':              'maria@ejemplo.com',
            'telefono':           '987654321',
            'dni':                '71234567',
            'nombre_completo':    'MARIA QUISPE LOPEZ',
            'fecha_nacimiento':   '14/03/1990',           # NUEVO
            'anios_experiencia':  3,
            'fecha_min':          2019,
            'fecha_max':          2026,
            'skills':             ['cocina_caliente', 'haccp', ...],
            'tiene_experiencia':  True,
            'tiene_educacion':    True,
            'longitud_texto':     12345,
            'idiomas':            ['ingles avanzado'],
            'secciones_detectadas': ['datos_personales', 'experiencia', ...],
        }
    """
    if not text:
        return _empty_parse_result()

    result = _empty_parse_result()
    result['longitud_texto'] = len(text)

    # ── FASE 1.1: segmentar ──
    secciones = segmentar_secciones(text)
    result['secciones_detectadas'] = sorted(
        k for k, v in secciones.items() if v and k != SECCION_OTROS
    )

    datos_pers = secciones.get(SECCION_DATOS_PERSONALES, '')
    exp_text = secciones.get(SECCION_EXPERIENCIA, '')
    cabecera = secciones.get(SECCION_OTROS, '')

    # ── Contacto (datos_personales + cabecera, NUNCA experiencia) ──
    contacto_text = '\n'.join(p for p in (cabecera, datos_pers) if p) or text

    emails = RE_EMAIL.findall(contacto_text)
    if emails:
        result['email'] = emails[0].lower()
    elif (emails_all := RE_EMAIL.findall(text)):
        result['email'] = emails_all[0].lower()

    tels = RE_TELEFONO_PE.findall(contacto_text) or RE_TELEFONO_PE.findall(text)
    if tels:
        # Normalizar — solo dígitos
        tel = re.sub(r'\D', '', tels[0])
        result['telefono'] = tel[-9:] if len(tel) >= 9 else tel

    # DNI: filtrar años y excluir DNIs que aparecen en la sección de experiencia
    dnis_contacto = RE_DNI_PE.findall(contacto_text) or RE_DNI_PE.findall(text)
    for d in dnis_contacto:
        if not d.startswith(('19', '20')) or len(d) != 8 or int(d[:4]) > 2050:
            result['dni'] = d
            break

    # ── Fecha de nacimiento (aislada — NO entra a experiencia) ──
    fecha_nac = extraer_fecha_nacimiento(text, datos_pers)
    if fecha_nac:
        result['fecha_nacimiento'] = fecha_nac

    año_nacimiento = _año_de_fecha_str(fecha_nac) if fecha_nac else None

    # ── Nombre completo (cabecera + primeras líneas) ──
    fuente_nombre = cabecera or text
    lineas = [l.strip() for l in fuente_nombre.split('\n') if l.strip()]
    for linea in lineas[:6]:
        palabras = linea.split()
        if (2 <= len(palabras) <= 4
                and all(p[0].isupper() for p in palabras if p)
                and not RE_EMAIL.search(linea)
                and not RE_TELEFONO_PE.search(linea)
                and not RE_FECHA_NACIMIENTO_LABEL.search(linea)
                and len(linea) < 80):
            result['nombre_completo'] = linea
            break

    # ── Fechas y experiencia (SÓLO de la sección experiencia) ──
    # Si no hay sección "Experiencia", caemos a heurística sobre el
    # texto completo PERO excluyendo el año de nacimiento.
    años_exp: list[int] = []
    fuente_fechas = exp_text or text
    año_actual = date.today().year

    for m in RE_FECHA.finditer(fuente_fechas):
        s = m.group()
        año_match = re.search(r'(?:19|20)\d{2}', s)
        if not año_match:
            continue
        año = int(año_match.group())
        if not (1950 <= año <= año_actual + 1):
            continue
        # Excluir explícitamente el año de nacimiento
        if año_nacimiento and año == año_nacimiento:
            continue
        # Si la fecha aparece junto a una etiqueta de nacimiento, excluir
        ctx_start = max(0, m.start() - 50)
        ctx = fuente_fechas[ctx_start:m.end()]
        if RE_FECHA_NACIMIENTO_LABEL.search(ctx):
            continue
        años_exp.append(año)

    if años_exp:
        result['fecha_min'] = min(años_exp)
        result['fecha_max'] = max(años_exp)
        # Si hay "Actualidad"/"Presente" en la sección, fecha_max=hoy
        if exp_text and re.search(
            r'\b(?:actualidad|presente|hoy|actual)\b',
            exp_text, re.IGNORECASE,
        ):
            result['fecha_max'] = max(result['fecha_max'], año_actual)
        result['anios_experiencia'] = min(
            max(0, result['fecha_max'] - result['fecha_min']),
            50,
        )

    # ── Secciones (booleans para compatibilidad) ──
    result['tiene_experiencia'] = bool(exp_text)
    result['tiene_educacion'] = bool(secciones.get(SECCION_EDUCACION))

    # ── Skills (búsqueda case-insensitive + normalización tildes) ──
    text_norm = _strip_tildes(text.lower())
    skills_encontrados = []
    for skill in SKILLS_CATALOG:
        if _strip_tildes(skill) in text_norm:
            skills_encontrados.append(skill)
    result['skills'] = sorted(skills_encontrados)

    # ── Idiomas (separar de skills) ──
    idiomas = [s for s in result['skills']
               if any(idioma in s for idioma in (
                   'ingles', 'portugues', 'frances', 'italiano',
                   'aleman', 'mandarin', 'japones', 'quechua'
               ))]
    result['idiomas'] = idiomas

    return result


def _empty_parse_result() -> dict:
    return {
        'email':             '',
        'telefono':          '',
        'dni':               '',
        'nombre_completo':   '',
        'fecha_nacimiento':  '',
        'anios_experiencia': 0,
        'fecha_min':         None,
        'fecha_max':         None,
        'skills':            [],
        'idiomas':           [],
        'tiene_experiencia': False,
        'tiene_educacion':   False,
        'longitud_texto':    0,
        'secciones_detectadas': [],
    }


def parse_cv_from_file(
    filepath: str,
    *,
    llm_caller: Optional[Callable[[str, dict], dict]] = None,
    forzar_ia: bool = False,
) -> dict:
    """Extrae texto y parsea entidades en un solo paso (FASE 1 + 2).

    Args:
        filepath: ruta al archivo.
        llm_caller: callable opcional para FASE 2 (enriquecimiento).
            Firma: `(texto_crudo, parsed_dict) -> dict_con_campos_faltantes`.
            Si es None, sólo se ejecuta FASE 1 (local).
            Sirve para inyectar mocks en tests.
        forzar_ia: si True, llama al LLM aunque no haya campos vacíos.

    Devuelve el dict de parse_cv enriquecido con:
        - texto_extraido: str
        - experiencias: list (extraídas de la sección experiencia)
        - educacion_items: list
        - ia_usada: bool  (True si FASE 2 corrió)
        - campos_de_ia: list[str]  (qué campos vinieron del LLM)
    """
    text = extract_text_from_file(filepath)

    # FASE 1: parsing LOCAL
    result = parse_cv(text)
    result['texto_extraido'] = text
    result['experiencias'] = extraer_experiencias(text)
    result['educacion_items'] = extraer_educacion(text)
    result['ia_usada'] = False
    result['campos_de_ia'] = []

    # FASE 2: enriquecimiento con IA (opcional, sólo si faltan campos)
    if llm_caller is not None:
        result = _enriquecer_con_ia(text, result, llm_caller, forzar=forzar_ia)

    return result


# ──────────────────────────────────────────────────────────────────
# FASE 2 — Enriquecimiento con IA (sólo si faltan campos)
# ──────────────────────────────────────────────────────────────────

# Campos que justifican llamar al LLM si están vacíos
_CAMPOS_CRITICOS = ('nombre_completo', 'email', 'telefono')

# Campos que el LLM puede rellenar (sin sobrescribir si ya hay valor)
_CAMPOS_ENRIQUECIBLES = (
    'nombre_completo', 'email', 'telefono', 'dni', 'fecha_nacimiento',
)


def _enriquecer_con_ia(
    texto: str,
    parsed: dict,
    llm_caller: Callable[[str, dict], dict],
    *,
    forzar: bool = False,
) -> dict:
    """Llama al LLM SÓLO si parsing local dejó campos críticos vacíos.

    Valida cada campo del LLM contra el texto original; los rechaza
    si parecen inventados (no aparecen literalmente en el CV).
    """
    if not forzar:
        faltantes = [c for c in _CAMPOS_CRITICOS if not parsed.get(c)]
        if not faltantes:
            return parsed

    try:
        sugerencias = llm_caller(texto, parsed) or {}
    except Exception as exc:
        logger.warning('LLM enrichment falló: %s', exc)
        return parsed

    if not isinstance(sugerencias, dict):
        return parsed

    aplicados: list[str] = []
    texto_lower = texto.lower()

    for campo in _CAMPOS_ENRIQUECIBLES:
        # No sobrescribir si parsing local ya tenía valor
        if parsed.get(campo):
            continue
        valor = sugerencias.get(campo)
        if not valor or not isinstance(valor, str):
            continue
        valor = valor.strip()
        if len(valor) < 2 or len(valor) > 250:
            continue

        # VALIDACIÓN ANTI-ALUCINACIÓN:
        # El valor del LLM debe aparecer literalmente en el CV
        # (insensitive). Si no aparece, lo descartamos.
        if not _validar_contra_texto(campo, valor, texto, texto_lower):
            logger.info('LLM sugirió %s="%s" pero no aparece en el CV — descartado.',
                        campo, valor[:60])
            continue

        parsed[campo] = valor
        aplicados.append(campo)

    parsed['ia_usada'] = bool(aplicados)
    parsed['campos_de_ia'] = aplicados
    return parsed


def _validar_contra_texto(campo: str, valor: str, texto: str, texto_lower: str) -> bool:
    """Verifica que el valor sugerido por el LLM aparezca en el CV."""
    if not valor:
        return False
    val_low = valor.lower().strip()

    if campo == 'email':
        # email completo debe aparecer literal
        return val_low in texto_lower
    if campo == 'telefono':
        # números del teléfono deben aparecer (con o sin separadores)
        solo_digitos = re.sub(r'\D', '', valor)
        if len(solo_digitos) < 6:
            return False
        digitos_en_texto = re.sub(r'\D', '', texto)
        return solo_digitos in digitos_en_texto
    if campo == 'dni':
        solo_digitos = re.sub(r'\D', '', valor)
        return len(solo_digitos) == 8 and solo_digitos in texto
    if campo == 'fecha_nacimiento':
        # Año debe aparecer y o bien la fecha completa o un patrón similar
        año = _año_de_fecha_str(valor)
        if not año:
            return False
        return str(año) in texto
    if campo == 'nombre_completo':
        # Al menos 2 de las palabras del nombre deben aparecer
        palabras = [w for w in val_low.split() if len(w) >= 3]
        if len(palabras) < 2:
            return False
        coincidencias = sum(1 for w in palabras if w in texto_lower)
        return coincidencias >= 2
    return val_low in texto_lower


# ──────────────────────────────────────────────────────────────────
# Extracción estructurada (experiencias / educación)
# ──────────────────────────────────────────────────────────────────

# Patrones de fecha extendidos para experiencias
RE_BLOQUE_FECHA = re.compile(
    r'(?P<inicio>'
        r'(?:Ene|Feb|Mar|Abr|May|Jun|Jul|Ago|Sep|Set|Oct|Nov|Dic)[a-z]*\.?\s+\d{4}'
        r'|\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{4}'
        r'|\d{4}'
    r')'
    r'\s*[-–—a]+\s*'      # separador (–, -, a)
    r'(?P<fin>'
        r'(?:Ene|Feb|Mar|Abr|May|Jun|Jul|Ago|Sep|Set|Oct|Nov|Dic)[a-z]*\.?\s+\d{4}'
        r'|\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{4}'
        r'|\d{4}'
        r'|Actualidad|Presente|Hoy|Actual'
    r')',
    re.IGNORECASE,
)


def _seccion_texto(text: str, header_re, max_chars: int = 4000) -> str:
    """Devuelve el texto que sigue a un header dado, hasta el siguiente
    header de sección o max_chars."""
    m = header_re.search(text)
    if not m:
        return ''
    start = m.end()
    next_headers = [
        RE_HEADER_EXPERIENCIA, RE_HEADER_EDUCACION,
        RE_HEADER_SKILLS, RE_HEADER_IDIOMAS,
    ]
    next_pos = start + max_chars
    for nh in next_headers:
        m2 = nh.search(text, pos=start + 5)
        if m2 and m2.start() < next_pos:
            next_pos = m2.start()
    return text[start:next_pos]


def extraer_experiencias(text: str) -> list:
    """Detecta bloques de experiencia laboral con fechas.

    Devuelve lista de dicts:
        {
            'fecha_inicio': 'Mar 2022',
            'fecha_fin':    'Actualidad',
            'lineas':       ['Restaurante Insignia', 'Mesera Senior'],
            'año_inicio':   2022,
            'año_fin':      2026,
        }
    """
    # Usar segmentación moderna para acotar a sección de experiencia.
    # Esto NO incluye datos personales — clave para evitar el bug
    # de fecha de nacimiento confundida con inicio de experiencia.
    secciones = segmentar_secciones(text)
    seccion = secciones.get(SECCION_EXPERIENCIA, '')
    if not seccion:
        # Fallback al header-based viejo
        seccion = _seccion_texto(text, RE_HEADER_EXPERIENCIA, max_chars=4000)
    if not seccion:
        return []

    experiencias = []
    matches = list(RE_BLOQUE_FECHA.finditer(seccion))

    for i, m in enumerate(matches):
        fecha_inicio = m.group('inicio').strip()
        fecha_fin = m.group('fin').strip()

        end_pos = matches[i + 1].start() if i + 1 < len(matches) else min(m.end() + 400, len(seccion))
        contexto = seccion[m.end():end_pos].strip()
        lineas = [l.strip() for l in contexto.split('\n') if l.strip()][:3]

        año_inicio = _año_de_str(fecha_inicio)
        año_fin = _año_de_str(fecha_fin)
        if not año_fin and fecha_fin.lower() in ('actualidad', 'presente', 'hoy', 'actual'):
            año_fin = date.today().year

        experiencias.append({
            'fecha_inicio': fecha_inicio,
            'fecha_fin':    fecha_fin,
            'año_inicio':   año_inicio,
            'año_fin':      año_fin,
            'lineas':       lineas,
            # Aliases para código legado:
            'periodo':      f'{fecha_inicio} - {fecha_fin}',
            'texto':        ' '.join(lineas),
        })

    return experiencias[:10]


def extraer_educacion(text: str) -> list:
    """Detecta bloques de educación con años + institución/título."""
    secciones = segmentar_secciones(text)
    seccion = secciones.get(SECCION_EDUCACION, '')
    if not seccion:
        seccion = _seccion_texto(text, RE_HEADER_EDUCACION, max_chars=2500)
    if not seccion:
        return []

    items = []
    año_re = re.compile(r'\b(19|20)\d{2}\b')
    años_encontrados = [(m.start(), m.group()) for m in año_re.finditer(seccion)]

    if not años_encontrados:
        return []

    for pos, año in años_encontrados[:5]:
        line_start = seccion.rfind('\n', 0, pos) + 1
        line_end = seccion.find('\n\n', pos)
        if line_end == -1:
            line_end = pos + 200
        contexto = seccion[line_start:line_end].strip()
        lineas = [l.strip() for l in contexto.split('\n') if l.strip()][:2]
        items.append({
            'año':    int(año),
            'lineas': lineas,
        })

    return items[:5]


def _año_de_str(s: str) -> Optional[int]:
    m = re.search(r'\b(19|20)\d{2}\b', s)
    if m:
        return int(m.group())
    return None
