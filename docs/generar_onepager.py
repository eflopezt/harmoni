"""
One-Pager / Ficha Tecnica de Harmoni ERP (1 sola pagina A4)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

TEAL = RGBColor(0x0F, 0x76, 0x6E)
DARK = RGBColor(0x1F, 0x2A, 0x37)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def shade(cell, color):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), color)
    sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

# === HEADER ===
t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
c0 = t.cell(0, 0)
c1 = t.cell(0, 1)
shade(c0, '0F766E')
shade(c1, '0F766E')

p0 = c0.paragraphs[0]
r = p0.add_run('H  HARMONI ERP')
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = WHITE

p1 = c1.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r1 = p1.add_run('ERP de RRHH y Planillas para Peru\nwww.harmoni.pe | +51 977 538 028')
r1.font.size = Pt(9)
r1.font.color.rgb = RGBColor(0xC8, 0xFA, 0xF2)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# === TAGLINE ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('22 modulos  |  100% normativa peruana  |  IA integrada  |  Desde S/149/mes')
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = TEAL

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# === 6 MODULOS GRID (3x2) ===
modules = [
    ('Personal & Contratos', 'Ficha de 50+ campos, organigrama interactivo,\ncontratos DL 728, PDF/DOCX, prorrogas, adendas'),
    ('Asistencia & Control', 'Calendario matricial, biometrico, papeletas,\nbanco de horas, reglas IA por persona'),
    ('Nominas & Planillas', 'Calculo automatico IR 5ta, AFP, CTS,\ngratificaciones, boletas PDF, liquidaciones'),
    ('Vacaciones DL 713', 'Saldos, solicitudes, triple vacacional,\nfraccionamiento DL 1405, calendario equipo'),
    ('Talento Humano', 'Reclutamiento Kanban, onboarding, evaluaciones\n360/OKR, capacitaciones, encuestas clima'),
    ('Integraciones Peru', 'PLAME, T-Registro, AFP Net, CTS bancos,\nEsSalud, CONCAR, SAP — en un click'),
]

t2 = doc.add_table(rows=3, cols=2)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (title, desc) in enumerate(modules):
    r, c = divmod(i, 2)
    cell = t2.cell(r, c)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    run_t = p.add_run(f'{title}\n')
    run_t.font.size = Pt(10)
    run_t.font.bold = True
    run_t.font.color.rgb = TEAL
    run_d = p.add_run(desc)
    run_d.font.size = Pt(8)
    run_d.font.color.rgb = GRAY

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# === DIFERENCIADORES ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Por que Harmoni')
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = DARK

diffs = [
    'IA Nativa: Asistente virtual, analisis contratos, reglas inteligentes, dashboard predictivo',
    'Multi-empresa: Row-level tenancy, SMTP por empresa, configuracion independiente',
    'Portal Trabajador: Autoservicio de boletas, vacaciones, papeletas, documentos',
    'PWA Movil: Funciona en cualquier dispositivo sin instalar aplicacion',
    '22 Modulos: Todo en una plataforma — no necesitas 5 sistemas diferentes',
    'Soporte Real: Ingenieros que entienden tu negocio, no un call center',
]

for d in diffs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    parts = d.split(': ')
    run1 = p.add_run(f'{parts[0]}: ')
    run1.font.size = Pt(8)
    run1.font.bold = True
    run1.font.color.rgb = TEAL
    run2 = p.add_run(parts[1])
    run2.font.size = Pt(8)
    run2.font.color.rgb = GRAY

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# === PRECIOS ===
t3 = doc.add_table(rows=2, cols=3)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['STARTER', 'PROFESIONAL', 'ENTERPRISE']
prices = ['S/ 149/mes\nHasta 50 colaboradores\nPersonal + Asistencia + Nominas',
          'S/ 349/mes\nHasta 200 colaboradores\n22 modulos + IA + Integraciones',
          'A medida\nColaboradores ilimitados\nMulti-empresa + API + SLA 24/7']

for i, h in enumerate(headers):
    cell = t3.cell(0, i)
    shade(cell, '0F766E' if i == 1 else '1F2A37')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = WHITE

for i, price in enumerate(prices):
    cell = t3.cell(1, i)
    if i == 1:
        shade(cell, 'E6FAF7')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(price)
    r.font.size = Pt(8)
    r.font.color.rgb = DARK if i == 1 else GRAY

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# === TECH STACK ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Stack: Django 5.1 + PostgreSQL 16 + Redis 7 + Celery 5 + Bootstrap 5 + Chart.js')
r.font.size = Pt(7)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
r.font.italic = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('IA Multi-proveedor: Gemini 2.5 Flash + DeepSeek V3 + GPT-4o-mini + pgvector (RAG)')
r2.font.size = Pt(7)
r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
r2.font.italic = True

# === FOOTER CTA ===
doc.add_paragraph().paragraph_format.space_after = Pt(2)
ft = doc.add_table(rows=1, cols=1)
ft.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = ft.cell(0, 0)
shade(cell, '0F766E')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('  Solicita tu DEMO GRATUITA: ventas@harmoni.pe  |  +51 977 538 028  |  harmoni.pe  ')
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = WHITE

out = r"D:\Harmoni\docs\FICHA TECNICA - HARMONI ERP.docx"
doc.save(out)
print(f"One-pager: {out}")

from docx2pdf import convert
pdf = out.replace('.docx', '.pdf')
convert(out, pdf)
print(f"PDF: {pdf}")
print("OK")
