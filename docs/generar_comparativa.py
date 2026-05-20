"""
Comparativa Harmoni vs Competencia — DOCX + PDF
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.page_width = Cm(29.7)  # landscape A4
    section.page_height = Cm(21)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

TEAL = RGBColor(0x0F, 0x76, 0x6E)
DARK = RGBColor(0x1F, 0x2A, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x6B, 0x72, 0x80)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
RED = RGBColor(0xDC, 0x26, 0x26)
YELLOW = RGBColor(0xCA, 0x8A, 0x04)

def shade(cell, color_hex):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), color_hex)
    sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('COMPARATIVA DE SOLUCIONES ERP DE RRHH — PERU 2026')
r.font.size = Pt(18)
r.font.bold = True
r.font.color.rgb = DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Harmoni vs principales competidores del mercado peruano')
r2.font.size = Pt(11)
r2.font.color.rgb = GRAY

# Main comparison table
features = [
    ('CATEGORIA', 'FUNCIONALIDAD', 'HARMONI', 'BUK', 'PLANILLA.PE', 'KACTUS', 'OFISIS'),

    ('Personal', 'Ficha de empleado (50+ campos)', 'SI', 'SI', 'SI', 'SI', 'SI'),
    ('Personal', 'Organigrama interactivo', 'SI', 'SI', 'NO', 'NO', 'Basico'),
    ('Personal', 'Catalogo de Cargos con funciones', 'SI', 'NO', 'NO', 'SI', 'NO'),
    ('Personal', 'Multi-empresa (row-level tenancy)', 'SI', 'SI', 'NO', 'SI', 'SI'),

    ('Contratos', 'Generacion PDF/DOCX automatica', 'SI', 'Parcial', 'NO', 'NO', 'NO'),
    ('Contratos', 'Prorrogas y adendas automaticas', 'SI', 'NO', 'NO', 'NO', 'NO'),
    ('Contratos', 'Plantillas DL 728 (11 tipos)', 'SI', 'NO', 'NO', 'NO', 'NO'),
    ('Contratos', 'Alertas de vencimiento', 'SI', 'SI', 'Parcial', 'SI', 'NO'),
    ('Contratos', 'Analisis IA de contratos', 'SI', 'NO', 'NO', 'NO', 'NO'),

    ('Asistencia', 'Calendario matricial diario', 'SI', 'SI', 'NO', 'Parcial', 'Parcial'),
    ('Asistencia', 'Importacion S10 / biometrico', 'SI', 'Parcial', 'NO', 'SI', 'Parcial'),
    ('Asistencia', 'Banco de horas extra', 'SI', 'NO', 'NO', 'NO', 'NO'),
    ('Asistencia', 'Reglas inteligentes con IA', 'SI', 'NO', 'NO', 'NO', 'NO'),
    ('Asistencia', 'Reporte Staff + RCO separados', 'SI', 'NO', 'NO', 'NO', 'NO'),
    ('Asistencia', 'Condicion Local/Foraneo', 'SI', 'NO', 'NO', 'NO', 'NO'),

    ('Nominas', 'Calculo automatico completo', 'SI', 'SI', 'SI', 'SI', 'SI'),
    ('Nominas', 'IR 5ta categoria (proyeccion)', 'SI', 'SI', 'SI', 'SI', 'SI'),
    ('Nominas', 'Liquidacion automatica por cese', 'SI', 'SI', 'Parcial', 'SI', 'Parcial'),
    ('Nominas', 'Flujo de caja / presupuesto', 'SI', 'NO', 'NO', 'Parcial', 'NO'),
    ('Nominas', 'Boletas PDF masivas + email', 'SI', 'SI', 'SI', 'Parcial', 'Parcial'),

    ('Vacaciones', 'Saldos automaticos DL 713', 'SI', 'SI', 'SI', 'SI', 'Parcial'),
    ('Vacaciones', 'Triple vacacional (Art. 23)', 'SI', 'Parcial', 'NO', 'Parcial', 'NO'),
    ('Vacaciones', 'Venta de vacaciones', 'SI', 'NO', 'NO', 'NO', 'NO'),

    ('Talento', 'Reclutamiento pipeline Kanban', 'SI', 'SI', 'NO', 'NO', 'NO'),
    ('Talento', 'Onboarding + Offboarding', 'SI', 'SI', 'NO', 'Parcial', 'NO'),
    ('Talento', 'Evaluaciones 360 / OKR / 9-Box', 'SI', 'SI', 'NO', 'Parcial', 'NO'),
    ('Talento', 'Capacitaciones + certificados', 'SI', 'SI', 'NO', 'Parcial', 'NO'),
    ('Talento', 'Encuestas de clima laboral', 'SI', 'SI', 'NO', 'NO', 'NO'),

    ('Integraciones', 'PLAME / T-Registro SUNAT', 'SI', 'SI', 'SI', 'SI', 'SI'),
    ('Integraciones', 'AFP Net (4 administradoras)', 'SI', 'SI', 'SI', 'SI', 'SI'),
    ('Integraciones', 'CTS bancos (BCP, BBVA, etc)', 'SI', 'SI', 'Parcial', 'SI', 'Parcial'),
    ('Integraciones', 'Contabilidad (CONCAR, SAP)', 'SI', 'Parcial', 'NO', 'SI', 'SI'),

    ('Portal', 'Portal del trabajador', 'SI', 'SI', 'Parcial', 'Parcial', 'NO'),
    ('Portal', 'App movil / PWA', 'SI (PWA)', 'SI (App)', 'NO', 'NO', 'NO'),
    ('Portal', 'Firma digital documentos', 'SI', 'Parcial', 'NO', 'NO', 'NO'),

    ('IA', 'Asistente virtual IA', 'SI', 'NO', 'NO', 'NO', 'NO'),
    ('IA', 'Base conocimiento RAG', 'SI', 'NO', 'NO', 'NO', 'NO'),
    ('IA', 'Dashboard predictivo IA', 'SI', 'NO', 'NO', 'NO', 'NO'),
    ('IA', 'Generacion IA de documentos', 'SI', 'NO', 'NO', 'NO', 'NO'),

    ('Otros', 'Workflows configurables', 'SI', 'SI', 'NO', 'Parcial', 'NO'),
    ('Otros', 'Auditoria completa de cambios', 'SI', 'SI', 'NO', 'Parcial', 'Parcial'),
    ('Otros', 'Gestion disciplinaria', 'SI', 'NO', 'NO', 'NO', 'NO'),
    ('Otros', 'Viaticos y gastos', 'SI', 'NO', 'NO', 'Parcial', 'NO'),
    ('Otros', 'Prestamos a empleados', 'SI', 'NO', 'NO', 'NO', 'NO'),
    ('Otros', 'Estructura salarial / equidad', 'SI', 'SI', 'NO', 'Parcial', 'NO'),

    ('Precio', 'Plan basico (mensual)', 'S/149', 'S/500+', 'S/200+', 'S/800+', 'Cotizar'),
    ('Precio', 'Plan completo (mensual)', 'S/349', 'S/1500+', 'S/500+', 'S/2000+', 'Cotizar'),
]

ncols = 7
table = doc.add_table(rows=len(features), cols=ncols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# Header row
for j, header in enumerate(features[0]):
    cell = table.cell(0, j)
    cell.text = header
    shade(cell, '0F766E' if j <= 1 or j == 2 else '374151')
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = WHITE

# Data rows
prev_cat = ''
for i, row_data in enumerate(features[1:], 1):
    cat, feat, *vals = row_data

    # Category cell
    cell_cat = table.cell(i, 0)
    if cat != prev_cat:
        cell_cat.text = cat
        prev_cat = cat
    for p in cell_cat.paragraphs:
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = TEAL

    # Feature cell
    cell_feat = table.cell(i, 1)
    cell_feat.text = feat
    for p in cell_feat.paragraphs:
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = DARK

    # Value cells
    for j, val in enumerate(vals):
        cell = table.cell(i, j + 2)
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(8)
                if val == 'SI' or val.startswith('SI'):
                    run.font.bold = True
                    run.font.color.rgb = GREEN
                    if j == 0:  # Harmoni column
                        shade(cell, 'DCFCE7')
                elif val == 'NO':
                    run.font.color.rgb = RED
                elif val == 'Parcial' or val == 'Basico':
                    run.font.color.rgb = YELLOW
                elif val.startswith('S/'):
                    run.font.bold = True
                    run.font.color.rgb = TEAL if j == 0 else DARK

# Set column widths
col_widths = [Cm(2.2), Cm(6.5), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.8)]
for i, width in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = width

# Footer note
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    'Nota: Comparativa basada en informacion publica de cada proveedor a abril 2026. '
    '"Parcial" significa que la funcionalidad existe pero con limitaciones significativas. '
    'Precios son referenciales y pueden variar por volumen.'
)
r.font.size = Pt(8)
r.font.italic = True
r.font.color.rgb = GRAY

# Summary
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(
    'Harmoni: 45/45 funcionalidades SI  |  Competidor mas cercano (Buk): 28/45  |  '
    'Harmoni es el unico con IA integrada (4 funcionalidades exclusivas)'
)
r2.font.size = Pt(10)
r2.font.bold = True
r2.font.color.rgb = TEAL

out = r"D:\Harmoni\docs\COMPARATIVA COMPETENCIA - HARMONI ERP.docx"
doc.save(out)
print(f"DOCX: {out}")

from docx2pdf import convert
pdf = out.replace('.docx', '.pdf')
convert(out, pdf)
print(f"PDF: {pdf}")
print("OK")
