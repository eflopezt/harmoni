"""
Generador del Manual de Usuario de Harmoni ERP
===============================================
Genera DOCX profesional con capturas, diagramas y descripciones.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\Harmoni\docs\manual_capturas"

doc = Document()

# ===== CONFIGURACION DE ESTILOS =====
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)  # teal
    h_style.font.name = 'Calibri'

# Helper functions
def add_image(path, width=Inches(6.5), caption=None):
    """Add image with optional caption."""
    full = os.path.join(OUT, path) if not os.path.isabs(path) else path
    if os.path.exists(full):
        doc.add_picture(full, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style.font.size = Pt(9)
            p.style.font.italic = True
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        doc.add_paragraph(f'[Imagen no disponible: {path}]')

def add_flow(steps, title=None):
    """Add a simple text-based flow diagram."""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    table = doc.add_table(rows=1, cols=len(steps))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, step in enumerate(steps):
        cell = table.cell(0, i)
        cell.text = step
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.bold = True
        # Shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '0F766E')
        shading.set(qn('w:val'), 'clear')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()

def add_bullet_list(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f'NOTA: {text}')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

def section_break():
    doc.add_page_break()


# ================================================================
# PORTADA
# ================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('HARMONI ERP')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema de Gestion de Recursos Humanos y Planillas')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MANUAL DE USUARIO')
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0.0  |  Abril 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Consorcio Stiler-Ripcon')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

section_break()

# ================================================================
# TABLA DE CONTENIDOS
# ================================================================
doc.add_heading('Tabla de Contenidos', level=1)
toc_items = [
    '1. Introduccion',
    '2. Acceso al Sistema',
    '3. Dashboard Principal',
    '4. Gestion de Personal',
    '   4.1 Empleados',
    '   4.2 Areas y Gerencias',
    '   4.3 Organigrama',
    '   4.4 Roster Matricial',
    '   4.5 Contratos Laborales',
    '   4.6 Reportes RRHH',
    '5. Control de Asistencia',
    '   5.1 Calendario Matricial',
    '   5.2 Importar Tareo',
    '   5.3 Papeletas de Permiso',
    '   5.4 Banco de Horas',
    '   5.5 Reporte Staff',
    '   5.6 Reporte RCO',
    '6. Nominas y Planillas',
    '   6.1 Periodos de Nomina',
    '   6.2 Boletas de Pago',
    '7. Vacaciones',
    '8. Reclutamiento',
    '9. Onboarding',
    '10. Workflows y Aprobaciones',
    '11. Documentos',
    '12. Comunicaciones',
    '13. Analytics',
    '14. Prestamos',
    '15. Integraciones y Exportaciones',
    '16. Cierre Mensual',
    '17. Portal del Trabajador',
    '18. Modulos Adicionales',
    '   18.1 Calendario Corporativo',
    '   18.2 Capacitaciones',
    '   18.3 Evaluaciones',
    '   18.4 Disciplinaria',
    '   18.5 Encuestas',
    '   18.6 Viaticos',
    '   18.7 Estructura Salarial',
    '19. Configuracion del Sistema',
    '20. Inteligencia Artificial',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('   '):
        for run in p.runs:
            run.font.bold = True

section_break()

# ================================================================
# 1. INTRODUCCION
# ================================================================
doc.add_heading('1. Introduccion', level=1)
doc.add_paragraph(
    'Harmoni es un sistema ERP integral de Recursos Humanos y Planillas disenado '
    'especificamente para empresas peruanas. Integra en una sola plataforma todos '
    'los procesos de gestion del talento humano: personal, asistencia, nominas, '
    'vacaciones, reclutamiento, evaluaciones y mas.'
)
doc.add_paragraph(
    'El sistema cumple al 100% con la normativa laboral peruana vigente: '
    'DL 728 (Ley de Productividad y Competitividad Laboral), DL 713 (Descansos '
    'Remunerados), DS 003-97-TR (LPCL), y todas las regulaciones de SUNAT, AFP y MTPE.'
)

doc.add_heading('Caracteristicas principales', level=2)
add_bullet_list([
    'Multi-empresa (row-level tenancy) con configuracion independiente por empresa',
    '22 modulos integrados que cubren todo el ciclo de vida del colaborador',
    'Normativa peruana al 100%: IR 5ta, AFP, ONP, gratificaciones, CTS, vacaciones',
    'Inteligencia Artificial integrada: asistente virtual, analisis de contratos, generacion de documentos',
    'Exportaciones oficiales: PLAME, T-Registro, AFP Net, CTS bancos',
    'Portal del Trabajador: autoservicio para consultas, solicitudes y documentos',
    'Reportes avanzados con graficos interactivos y exportacion a Excel/PDF',
    'PWA (Progressive Web App): acceso desde moviles sin instalar',
])

section_break()

# ================================================================
# 2. ACCESO AL SISTEMA
# ================================================================
doc.add_heading('2. Acceso al Sistema', level=1)
doc.add_paragraph(
    'Para ingresar a Harmoni, abra su navegador web y acceda a la URL proporcionada '
    'por su administrador. Se mostrara la pantalla de inicio de sesion.'
)

add_image('00_login/01_login_page.png', caption='Figura 2.1 - Pantalla de inicio de sesion')

doc.add_paragraph(
    'Ingrese su usuario y contrasena. El sistema valida las credenciales y redirige '
    'al Dashboard principal. Si olvido su contrasena, haga click en "Olvide mi contrasena" '
    'para recibir un enlace de recuperacion por email.'
)

add_flow(['Abrir URL', 'Ingresar credenciales', 'Validar acceso', 'Dashboard'],
         'Flujo de acceso:')

section_break()

# ================================================================
# 3. DASHBOARD
# ================================================================
doc.add_heading('3. Dashboard Principal', level=1)
doc.add_paragraph(
    'El Dashboard es la pantalla central de Harmoni. Muestra un resumen ejecutivo '
    'de toda la gestion de RRHH en tiempo real.'
)

add_image('01_dashboard/01_main.png', caption='Figura 3.1 - Dashboard principal con KPIs')

doc.add_heading('Elementos del Dashboard', level=2)
add_bullet_list([
    'Saludo personalizado con fecha y conteo de colaboradores activos',
    'Frase motivacional del dia (cambia automaticamente)',
    'Grafico donut: distribucion Staff vs RCO con total de colaboradores',
    'KPIs de asistencia del dia: Presentes, Faltas, Permisos',
    'Alerta de contratos por vencer en los proximos 30 dias',
    'Tarjetas de acceso rapido a todos los modulos del sistema',
    'Seccion Portal del Trabajador con accesos personalizados',
    'Indicadores de periodos de prueba y cumpleanos del dia',
])

add_image('01_dashboard/02_accesos_rapidos.png', caption='Figura 3.2 - Accesos rapidos y portal del trabajador')

section_break()

# ================================================================
# 4. GESTION DE PERSONAL
# ================================================================
doc.add_heading('4. Gestion de Personal', level=1)
doc.add_paragraph(
    'El modulo de Personal es el nucleo de Harmoni. Gestiona toda la informacion '
    'de los colaboradores desde su ingreso hasta su cese.'
)

# 4.1 Empleados
doc.add_heading('4.1 Lista de Empleados', level=2)
add_image('02_personal/01_lista_empleados.png', caption='Figura 4.1 - Lista de empleados con filtros')
doc.add_paragraph(
    'La lista de empleados muestra todos los colaboradores con filtros por estado '
    '(activo/cesado), area, proyecto y condicion (Staff/RCO). La busqueda permite '
    'encontrar por nombre o DNI.'
)

doc.add_heading('Ficha del Empleado', level=3)
add_image('02_personal/02_detalle_general.png', caption='Figura 4.2 - Ficha del empleado - datos generales')
doc.add_paragraph(
    'La ficha individual contiene: datos personales, informacion laboral (cargo, area, '
    'proyecto, regimen de turno), datos bancarios, informacion tributaria (AFP/ONP), '
    'historial de contratos, documentos adjuntos y configuraciones especiales de asistencia.'
)

add_image('02_personal/03_detalle_scroll1.png', caption='Figura 4.3 - Ficha del empleado - seccion inferior')

doc.add_heading('Registrar Nuevo Empleado', level=3)
add_image('02_personal/12_nuevo_empleado.png', caption='Figura 4.4 - Formulario de alta de nuevo empleado')

add_flow(['Nuevo Empleado', 'Datos personales', 'Info laboral', 'Contrato', 'Onboarding'],
         'Flujo de alta de personal:')

# 4.2 Areas
doc.add_heading('4.2 Areas y Gerencias', level=2)
add_image('02_personal/04_areas.png', caption='Figura 4.5 - Estructura de areas y gerencias')
doc.add_paragraph(
    'Harmoni organiza la empresa en Areas/Gerencias y Sub-Areas. Cada area tiene '
    'un gerente responsable y contabiliza automaticamente el numero de colaboradores.'
)
add_image('02_personal/05_subareas.png', caption='Figura 4.6 - Sub-areas')

# 4.3 Organigrama
doc.add_heading('4.3 Organigrama Corporativo', level=2)
add_image('02_personal/06_organigrama.png', caption='Figura 4.7 - Organigrama interactivo')
doc.add_paragraph(
    'El organigrama muestra la estructura jerarquica de la empresa con nodos coloreados '
    'por nivel: Alta Direccion (esmeralda), Gerencia (verde), Jefatura (azul), Coordinacion '
    '(celeste), Especialista/Tecnico (amarillo), Operativo (gris). Es interactivo con '
    'zoom, drag, y exportacion a PDF/imagen.'
)

# 4.4 Roster
doc.add_heading('4.4 Roster Matricial', level=2)
add_image('02_personal/07_roster.png', caption='Figura 4.8 - Roster matricial')

# 4.5 Contratos
doc.add_heading('4.5 Contratos Laborales', level=2)
add_image('02_personal/09_contratos_panel.png', caption='Figura 4.9 - Panel de contratos laborales')
doc.add_paragraph(
    'El modulo de contratos gestiona todo el ciclo contractual segun el DL 728. '
    'Muestra KPIs: contratos activos, indefinidos, plazo fijo, vencidos, por vencer '
    'en 30 dias. Pestanas: Vencimientos, Periodo de Prueba, Por Modalidad, Actividad Reciente.'
)

add_image('02_personal/10_contratos_lista.png', caption='Figura 4.10 - Lista completa de contratos')

doc.add_heading('Detalle de Contrato', level=3)
add_image('02_personal/10_contrato_detalle.png', caption='Figura 4.11 - Detalle de contrato individual')
doc.add_paragraph(
    'El detalle de cada contrato muestra: fecha ingreso, modalidad, estado, periodo de prueba. '
    'Acciones rapidas: Generar PDF, Descargar DOCX, Generar Prorroga, Adenda (Aumento Salarial, '
    'Cambio de Cargo, Cambio de Condiciones), Enviar por Email. Los documentos se generan '
    'automaticamente con plantillas basadas en la normativa DL 728.'
)

add_flow(['Crear contrato', 'Asignar plantilla', 'Generar PDF/DOCX', 'Prorroga/Adenda', 'Enviar email'],
         'Flujo de gestion contractual:')

# 4.6 Reportes
doc.add_heading('4.6 Reportes de RRHH', level=2)
add_image('02_personal/11_reportes.png', caption='Figura 4.12 - Centro de reportes de RRHH')

section_break()

# ================================================================
# 5. ASISTENCIA
# ================================================================
doc.add_heading('5. Control de Asistencia', level=1)
doc.add_paragraph(
    'El modulo de Asistencia controla la jornada laboral de todos los colaboradores. '
    'Soporta multiples regimenes de turno (5x2, 14x7, 21x7, etc.), jornadas diferenciadas '
    'para personal Local y Foraneo, y reglas especiales configurables por persona.'
)

add_image('03_asistencia/01_dashboard.png', caption='Figura 5.1 - Dashboard de asistencia')

# 5.1 Calendario
doc.add_heading('5.1 Calendario Matricial', level=2)
add_image('03_asistencia/03_calendario_marzo.png', caption='Figura 5.2 - Asistencia Matricial - Marzo 2026')
doc.add_paragraph(
    'La Asistencia Matricial es la vista principal: grilla diaria por empleado para todo '
    'el mes. KPIs superiores: total empleados, dias trabajados, faltas, tasa de asistencia. '
    'Codigos de colores: A (Asistio/verde), DS (Descanso Semanal/gris), DL (Dia Libre), '
    'VAC (Vacaciones/celeste), SS (Sin Salida/amarillo), CHE (Compensacion HE), '
    'NA (No Aplica), FER (Feriado). Click en cualquier celda muestra el detalle de marcaciones.'
)

doc.add_heading('Filtros disponibles', level=3)
add_bullet_list([
    'Mes y anio',
    'Estado: Todos, Solo activos, Solo cesados',
    'Area: filtro por area/gerencia',
    'Condicion: Local, Foraneo, Todos',
    'Busqueda por nombre o DNI',
    'Vista: Mes completo (1-31) o Ciclo (21-20)',
])

# 5.2 Importar
doc.add_heading('5.2 Importar Tareo / Biometrico', level=2)
add_image('03_asistencia/04_importar.png', caption='Figura 5.3 - Importar registros de asistencia')
doc.add_paragraph(
    'Harmoni importa asistencia desde multiples fuentes: Excel S10 (formato detalle y resumen), '
    'biometrico (diversas marcas), e importacion manual. El sistema detecta automaticamente '
    'inconsistencias: marcas sin ingreso, jornadas incompletas, y aplica las reglas del regimen.'
)

add_flow(['Subir archivo', 'Validar formato', 'Procesar marcas', 'Aplicar reglas', 'Generar registros'],
         'Flujo de importacion:')

# 5.3 Papeletas
doc.add_heading('5.3 Papeletas de Permiso', level=2)
add_image('03_asistencia/05_papeletas.png', caption='Figura 5.4 - Gestion de papeletas')

# 5.4 Banco
doc.add_heading('5.4 Banco de Horas Extra', level=2)
add_image('03_asistencia/06_banco_horas.png', caption='Figura 5.5 - Banco de horas extra')
doc.add_paragraph(
    'El Banco de Horas acumula las horas extra trabajadas. El trabajador puede compensar '
    'las horas acumuladas como descanso en otro momento. Se integra con el calendario '
    'de asistencia para reflejar las compensaciones (CHE).'
)

# 5.5 Staff
doc.add_heading('5.5 Reporte de Asistencia Staff', level=2)
add_image('03_asistencia/07_reporte_staff.png', caption='Figura 5.6 - Reporte Staff - Marzo 2026')
doc.add_paragraph(
    'Reporte detallado para personal administrativo (Staff): 161 personas activas, '
    '4,202.8 horas extra acumuladas, 1,652 horas en banco disponibles, 3,317 dias '
    'trabajados, 51 faltas. Columnas: Trabajado, SS, DL, CHE, VAC, AUS, HE25%, HE35%, '
    'HE100%, Saldo Banco, Compensaciones.'
)

# 5.6 RCO
doc.add_heading('5.6 Reporte RCO - Regimen Civil de Obra', level=2)
add_image('03_asistencia/08_reporte_rco.png', caption='Figura 5.7 - Reporte RCO - Marzo 2026')
doc.add_paragraph(
    'Reporte especifico para personal de obra (RCO): 49 personas, desglose de horas extra '
    'por tipo (HE 25%: 1,766h, HE 35%: 819h, HE 100%: 484.4h). Total 3,069.4 horas extra '
    'en 1,403 registros. Vista Resumen y Detalle Diario. Botones: Exportar Horas, Carga S10.'
)

section_break()

# ================================================================
# 6. NOMINAS
# ================================================================
doc.add_heading('6. Nominas y Planillas', level=1)
doc.add_paragraph(
    'El motor de nominas de Harmoni calcula automaticamente todos los conceptos '
    'remunerativos y deductivos segun la normativa peruana: sueldo basico, '
    'asignacion familiar, horas extra, gratificaciones, CTS, AFP/ONP, IR 5ta categoria, '
    'prestamos, adelantos y mas.'
)

add_image('04_nominas/01_periodos.png', caption='Figura 6.1 - Lista de periodos de nomina')
add_image('04_nominas/02_periodo_detalle.png', caption='Figura 6.2 - Detalle de periodo')
add_image('04_nominas/03_periodo_tabla.png', caption='Figura 6.3 - Tabla de nomina por empleado')

doc.add_heading('6.2 Boletas de Pago', level=2)
add_image('04_nominas/04_boletas.png', caption='Figura 6.4 - Generacion de boletas PDF')

add_flow(['Abrir periodo', 'Calcular nomina', 'Revisar', 'Generar boletas', 'Enviar email', 'Cerrar periodo'],
         'Flujo de nomina mensual:')

section_break()

# ================================================================
# 7. VACACIONES
# ================================================================
doc.add_heading('7. Vacaciones', level=1)
doc.add_paragraph(
    'Gestion completa de vacaciones segun DL 713: 30 dias por ano de servicio. '
    'Control de saldos, solicitudes con flujo de aprobacion, vacaciones truncas, '
    'vencidas (triple vacacional Art. 23) y fraccionamiento (DL 1405).'
)
add_image('05_vacaciones/01_panel.png', caption='Figura 7.1 - Panel de vacaciones')
add_image('05_vacaciones/02_saldos.png', caption='Figura 7.2 - Saldos de vacaciones por trabajador')

add_flow(['Solicitar', 'Aprobacion jefe', 'Programar', 'Gozar', 'Registrar'],
         'Flujo de solicitud de vacaciones:')

section_break()

# ================================================================
# 8-12. MODULOS ADICIONALES (paginas individuales)
# ================================================================
modules = [
    ('8. Reclutamiento', '06_reclutamiento/01_vacantes.png', 'Figura 8.1',
     'Gestion de vacantes y procesos de seleccion: publicacion, recepcion de CVs, evaluacion, entrevistas, seleccion y contratacion.'),
    ('9. Onboarding', '07_onboarding/01_panel.png', 'Figura 9.1',
     'Proceso de incorporacion del nuevo colaborador: checklist de documentos, induccion, asignacion de recursos, firma de contrato, presentacion al equipo.'),
    ('10. Workflows y Aprobaciones', '08_workflows/01_aprobaciones.png', 'Figura 10.1',
     'Flujos de aprobacion configurables: vacaciones, permisos, prestamos, gastos. Escalamiento automatico por tiempo. Notificaciones por email.'),
    ('11. Gestion Documental', '09_documentos/01_panel.png', 'Figura 11.1',
     'Repositorio centralizado de documentos por empleado: contratos, DNI, certificados, constancias, memorandums. Control de vencimientos y descarga.'),
    ('12. Comunicaciones Internas', '10_comunicaciones/01_panel.png', 'Figura 12.1',
     'Sistema de comunicaciones: anuncios generales, circulares por area, memorandums individuales, notificaciones push. Historial de envios y confirmacion de lectura.'),
]

for title, img, caption, desc in modules:
    doc.add_heading(title, level=1)
    doc.add_paragraph(desc)
    add_image(img, caption=caption + ' - ' + title)
    doc.add_paragraph()

section_break()

# ================================================================
# 13. ANALYTICS
# ================================================================
doc.add_heading('13. Analytics', level=1)
doc.add_paragraph(
    'Dashboard de metricas avanzadas con graficos interactivos: headcount por area, '
    'evolucion mensual, indice de rotacion, costo laboral, distribucion por cargo, '
    'genero, edad y antiguedad. Exportable a Excel e imagen.'
)
add_image('11_analytics/01_dashboard.png', caption='Figura 13.1 - Dashboard de Analytics')
add_image('11_analytics/02_graficos.png', caption='Figura 13.2 - Graficos detallados')

section_break()

# ================================================================
# 14-16
# ================================================================
doc.add_heading('14. Prestamos', level=1)
doc.add_paragraph(
    'Control de prestamos a empleados: monto, cuotas, tasa, estado de pago. '
    'Descuento automatico en nomina. Simulador de cuotas.'
)
add_image('12_prestamos/01_panel.png', caption='Figura 14.1 - Gestion de prestamos')

doc.add_heading('15. Integraciones y Exportaciones', level=1)
doc.add_paragraph(
    'Exportaciones a entidades oficiales peruanas: PLAME mensual (SUNAT), T-Registro '
    'altas/bajas, AFP Net (aportes previsionales), CTS bancos (BCP Telecredito, BBVA, '
    'Interbank), PDT 601. Cada exportacion genera archivo en formato oficial vigente.'
)
add_image('13_integraciones/01_panel.png', caption='Figura 15.1 - Centro de integraciones')

doc.add_heading('16. Cierre Mensual', level=1)
doc.add_paragraph(
    'Proceso de cierre que valida la completitud de datos del periodo: asistencia '
    'procesada, nomina calculada, documentos generados. Bloquea edicion retroactiva '
    'para garantizar integridad de datos.'
)
add_image('14_cierre/01_panel.png', caption='Figura 16.1 - Panel de cierre mensual')

section_break()

# ================================================================
# 17. PORTAL DEL TRABAJADOR
# ================================================================
doc.add_heading('17. Portal del Trabajador', level=1)
doc.add_paragraph(
    'El Portal del Trabajador es la interfaz de autoservicio para los colaboradores. '
    'Cada empleado accede con sus credenciales personales y puede consultar su '
    'informacion, descargar boletas, solicitar vacaciones y permisos, y ver comunicaciones.'
)

add_image('15_portal/01_portal_main.png', caption='Figura 17.1 - Inicio del Portal del Trabajador')

doc.add_heading('Funcionalidades del Portal', level=2)

portal_items = [
    ('15_portal/02_mis_datos.png', 'Figura 17.2 - Mis Datos Personales',
     'Consulta de datos personales, solicitud de actualizacion.'),
    ('15_portal/03_mi_asistencia.png', 'Figura 17.3 - Mi Asistencia',
     'Marcaciones de entrada/salida, calendario personal, horas trabajadas.'),
    ('15_portal/04_mis_vacaciones.png', 'Figura 17.4 - Mis Vacaciones',
     'Saldo de vacaciones, solicitar nuevas vacaciones, historial de goce.'),
    ('15_portal/05_mis_papeletas.png', 'Figura 17.5 - Mis Papeletas',
     'Crear solicitudes de permiso, ver estado de aprobacion, historial.'),
]

for img, caption, desc in portal_items:
    doc.add_paragraph(desc)
    add_image(img, caption=caption)

section_break()

# ================================================================
# 18. MODULOS ADICIONALES
# ================================================================
doc.add_heading('18. Modulos Adicionales', level=1)

extras = [
    ('18.1 Calendario Corporativo', '16_calendario/01_calendario.png',
     'Calendario centralizado con feriados, vacaciones del equipo, cumpleanos y eventos.'),
    ('18.2 Capacitaciones', '17_capacitaciones/01_panel.png',
     'Plan de capacitaciones: cursos, talleres, certificaciones. Control de asistencia y evaluaciones.'),
    ('18.3 Evaluaciones de Desempeno', '18_evaluaciones/01_panel.png',
     'Evaluaciones por objetivos (OKR), competencias, evaluacion 360. Ciclos configurables.'),
    ('18.4 Gestion Disciplinaria', '19_disciplinaria/01_panel.png',
     'Registro de medidas disciplinarias: amonestaciones, suspensiones. Base legal, historial.'),
    ('18.5 Encuestas de Clima', '20_encuestas/01_panel.png',
     'Encuestas de clima laboral y satisfaccion: formularios, envio masivo, resultados graficos.'),
    ('18.6 Viaticos', '21_viaticos/01_panel.png',
     'Solicitudes de viaticos, rendiciones, liquidaciones. Flujo de aprobacion e integracion con nomina.'),
    ('18.7 Estructura Salarial', '22_salarios/01_panel.png',
     'Bandas salariales por cargo, historico de incrementos, analisis de equidad interna.'),
]

for title, img, desc in extras:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)
    add_image(img)
    doc.add_paragraph()

section_break()

# ================================================================
# 19. CONFIGURACION
# ================================================================
doc.add_heading('19. Configuracion del Sistema', level=1)

add_image('23_empresas/01_panel.png', caption='Figura 19.1 - Configuracion de empresa')
doc.add_paragraph(
    'Configuracion multi-empresa: razon social, RUC, representante legal (tipo documento, '
    'numero), logo y firma digital para documentos. Configuracion SMTP por empresa para '
    'envio de emails (Gmail, Office 365, custom). Row-level tenancy multi-tenant.'
)

add_image('24_sistema/01_panel.png', caption='Figura 19.2 - Panel de sistema')

section_break()

# ================================================================
# 20. IA
# ================================================================
doc.add_heading('20. Inteligencia Artificial', level=1)
doc.add_paragraph(
    'Harmoni integra IA de multiples proveedores (Gemini, DeepSeek, GPT) para '
    'potenciar la gestion de RRHH:'
)

add_bullet_list([
    'Asistente virtual: responde consultas sobre personal, asistencia, normativa laboral',
    'Analisis de contratos: revisa clausulas y sugiere mejoras segun DL 728',
    'Generacion de documentos: contratos, adendas, prorrogas con plantillas inteligentes',
    'Base de conocimiento (RAG): 59 articulos de normativa peruana indexados con pgvector',
    'Dashboard inteligente: frases motivacionales, alertas proactivas, resumen ejecutivo',
    'Busqueda universal: Ctrl+K para buscar empleado, papeleta, documento en todo el sistema',
])

doc.add_paragraph(
    'El boton flotante del chatbot (esquina inferior derecha) permite acceder al '
    'asistente IA desde cualquier pantalla del sistema.'
)

section_break()

# ================================================================
# LANDING PAGE
# ================================================================
doc.add_heading('Landing Page', level=1)
add_image('25_landing/01_landing_hero.png', caption='Landing page publica de Harmoni')
add_image('25_landing/02_landing_full.png', caption='Vista completa de la landing page')

# ================================================================
# GUARDAR
# ================================================================
out_path = r"D:\Harmoni\docs\MANUAL DE USUARIO - HARMONI ERP.docx"
doc.save(out_path)
print(f"Guardado: {out_path}")
print("Paginas estimadas: ~70+")
print("COMPLETADO")
