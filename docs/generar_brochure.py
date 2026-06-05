"""
Generador del Brochure Comercial de Harmoni ERP
================================================
PDF profesional para clientes potenciales.
Colores corporativos: Teal #0F766E, Dark #0A2F2B, Cyan #5EEAD4
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMGS = r"D:\Harmoni\docs\manual_capturas"

doc = Document()

# Page setup - A4
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# Colors
TEAL = RGBColor(0x0F, 0x76, 0x6E)
DARK = RGBColor(0x0A, 0x2F, 0x2B)
CYAN = RGBColor(0x5E, 0xEA, 0xD4)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x6B, 0x72, 0x80)
DARKGRAY = RGBColor(0x33, 0x33, 0x33)

def add_centered(text, size=12, color=DARKGRAY, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    return p

def add_left(text, size=10, color=DARKGRAY, bold=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    return p

def add_icon_text(icon, text, size=10, color=DARKGRAY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'{icon}  {text}')
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_img(path, width=Inches(6.2), caption=None):
    full = os.path.join(IMGS, path) if not os.path.isabs(path) else path
    if os.path.exists(full):
        doc.add_picture(full, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = GRAY

def add_divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 80)
    run.font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)
    run.font.size = Pt(6)

def shade_cell(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def feature_table(features):
    """Create a 2-column feature grid."""
    rows = (len(features) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, feat in enumerate(features):
        r, c = divmod(i, 2)
        cell = table.cell(r, c)
        cell.text = feat
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = DARKGRAY
    doc.add_paragraph()

# ================================================================
# PAGE 1 - PORTADA
# ================================================================
for _ in range(5):
    doc.add_paragraph()

add_centered('H', size=72, color=TEAL, bold=True, space_after=0)
add_centered('HARMONI', size=42, color=TEAL, bold=True, space_after=4)
add_centered('ERP de Recursos Humanos y Planillas', size=16, color=GRAY, space_after=20)

add_divider()
doc.add_paragraph()

add_centered('El sistema integral que tu empresa necesita', size=20, color=DARK, bold=True, space_after=8)
add_centered('para gestionar todo el ciclo de vida del colaborador.', size=14, color=GRAY, space_after=20)

add_centered('Normativa Peruana al 100%  |  22 Modulos Integrados  |  IA Incluida', size=11, color=TEAL, bold=True, space_after=20)

doc.add_paragraph()
doc.add_paragraph()
add_centered('www.harmoni.pe', size=14, color=TEAL, bold=True, space_after=4)
add_centered('+51 977 538 028  |  ventas@harmoni.pe', size=11, color=GRAY)

doc.add_page_break()

# ================================================================
# PAGE 2 - PROBLEMA / SOLUCION
# ================================================================
add_centered('El problema que resolvemos', size=22, color=DARK, bold=True, space_after=12)

problemas = [
    'Tu planilla se calcula en Excel con riesgo de errores y multas SUNAFIL',
    'Los contratos se hacen a mano y se vencen sin que nadie se entere',
    'No tienes control real de asistencia: las horas extra se pierden',
    'Las vacaciones se calculan "a ojo" y no cumples con el DL 713',
    'Cada exportacion a SUNAT/AFP/bancos toma horas de trabajo manual',
    'No tienes visibilidad en tiempo real de tu fuerza laboral',
]
for prob in problemas:
    p = doc.add_paragraph()
    run = p.add_run(f'   {prob}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.add_paragraph()
add_divider()
doc.add_paragraph()

add_centered('Harmoni lo resuelve todo en una sola plataforma', size=18, color=TEAL, bold=True, space_after=12)

soluciones = [
    'Motor de nominas que calcula automaticamente IR 5ta, AFP, gratificaciones, CTS',
    'Gestion de contratos con alertas, generacion PDF/DOCX y cumplimiento DL 728',
    'Control biometrico de asistencia con reglas inteligentes por persona',
    'Vacaciones calculadas al dia con saldos, vencidas y triple vacacional',
    'Exportacion directa a PLAME, T-Registro, AFP Net, CTS bancos en un click',
    'Dashboard en tiempo real con 200+ indicadores y graficos interactivos',
]
for sol in soluciones:
    p = doc.add_paragraph()
    run = p.add_run(f'   {sol}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

doc.add_page_break()

# ================================================================
# PAGE 3 - SCREENSHOT DASHBOARD
# ================================================================
add_centered('Vista del Dashboard Principal', size=18, color=DARK, bold=True, space_after=8)
add_centered('Todo tu equipo en una sola pantalla', size=12, color=GRAY, space_after=12)
add_img('01_dashboard/01_main.png', width=Inches(6.5))
doc.add_paragraph()

# KPI boxes
kpis = [
    ('200+', 'Colaboradores\nen tiempo real'),
    ('22', 'Modulos\nintegrados'),
    ('96.4%', 'Tasa asistencia\nautomatica'),
    ('100%', 'Normativa\nperuana'),
]
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (num, label) in enumerate(kpis):
    cell = table.cell(0, i)
    shade_cell(cell, '0F766E')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{num}\n')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run2 = p.add_run(label)
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0xC8, 0xFA, 0xF2)

doc.add_page_break()

# ================================================================
# PAGE 4 - MODULOS
# ================================================================
add_centered('22 Modulos que cubren todo el ciclo del colaborador', size=18, color=DARK, bold=True, space_after=16)

modulos = [
    ('GESTION DE PERSONAL', [
        'Empleados — Ficha completa con 50+ campos',
        'Organigrama — Interactivo con colores por nivel',
        'Contratos — DL 728, PDF/DOCX, prorrogas, adendas',
        'Cargos — Catalogo normalizado con funciones',
        'Areas/SubAreas — Estructura organizacional',
        'Roster Matricial — Planificacion de personal',
    ]),
    ('ASISTENCIA Y CONTROL', [
        'Calendario Matricial — Grilla diaria por empleado',
        'Importacion — S10, biometrico, manual',
        'Papeletas — Permisos con flujo aprobacion',
        'Banco de Horas — Acumulacion y compensacion HE',
        'Reporte Staff — Horas normales, extras, nocturnas',
        'Reporte RCO — Regimen Civil de Obra',
    ]),
    ('NOMINAS Y COMPENSACIONES', [
        'Planilla — Calculo automatico todos los conceptos',
        'Boletas PDF — Individual y masiva',
        'Liquidaciones — Calculo automatico por cese',
        'IR 5ta — Proyeccion y calculo mensual',
        'Flujo de Caja — Presupuesto vs ejecutado',
        'Prestamos — Control de cuotas y descuentos',
    ]),
    ('TALENTO HUMANO', [
        'Reclutamiento — Pipeline Kanban de candidatos',
        'Onboarding — Checklist de incorporacion',
        'Evaluaciones — 360, OKR, 9-Box, competencias',
        'Capacitaciones — Plan anual, certificados',
        'Encuestas — Clima laboral, satisfaccion',
        'Disciplinaria — Amonestaciones, base legal',
    ]),
    ('OPERACIONES Y COMPLIANCE', [
        'Vacaciones — DL 713, saldos, triple vacacional',
        'Documentos — Legajo digital, firma electronica',
        'Comunicaciones — Anuncios, WhatsApp, notificaciones',
        'Workflows — Aprobaciones multi-nivel configurables',
        'Integraciones — PLAME, T-Registro, AFP, CTS bancos',
        'Cierre — Validaciones y bloqueo de periodo',
    ]),
]

for grupo, items in modulos:
    p = doc.add_paragraph()
    run = p.add_run(grupo)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = TEAL
    p.paragraph_format.space_after = Pt(4)

    for item in items:
        parts = item.split(' — ')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f'  {parts[0]}')
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = DARKGRAY
        if len(parts) > 1:
            run2 = p.add_run(f' — {parts[1]}')
            run2.font.size = Pt(9)
            run2.font.color.rgb = GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ================================================================
# PAGE 5 - ASISTENCIA SCREENSHOT
# ================================================================
add_centered('Control de Asistencia en Tiempo Real', size=18, color=DARK, bold=True, space_after=8)
add_img('03_asistencia/03_calendario_marzo.png', width=Inches(6.5), caption='Asistencia Matricial — Marzo 2026')
doc.add_paragraph()
add_left(
    'Grilla diaria por empleado con codigos de color: Asistio (verde), Descanso (gris), '
    'Vacaciones (celeste), Falta (rojo), Licencia (turquesa). 207 empleados procesados '
    'automaticamente con 96.4% de tasa de asistencia. Click en cualquier celda para ver '
    'detalle de marcaciones.',
    size=10, color=GRAY
)

doc.add_paragraph()
add_centered('Reporte Staff con Horas Extra', size=14, color=DARK, bold=True, space_after=8)
add_img('03_asistencia/07_reporte_staff.png', width=Inches(6.5), caption='161 personas Staff — 4,202 horas extra acumuladas')

doc.add_page_break()

# ================================================================
# PAGE 6 - CONTRATOS
# ================================================================
add_centered('Gestion de Contratos Laborales', size=18, color=DARK, bold=True, space_after=8)
add_centered('Cumplimiento total del DL 728 / DS 003-97-TR', size=12, color=GRAY, space_after=12)
add_img('02_personal/09_contratos_panel.png', width=Inches(6.5), caption='Panel de contratos con alertas de vencimiento')
doc.add_paragraph()

contratos_features = [
    '11 plantillas basadas en normativa real (Confianza, Fiscalizable, Obra Determinada)',
    'Generacion automatica de PDF y DOCX editables',
    'Prorrogas y adendas (aumento salarial, cambio de cargo, cambio de condiciones)',
    'Alertas: 57 contratos vencen en los proximos 30 dias',
    'Envio masivo por email con un click',
    'Analisis de contratos con Inteligencia Artificial',
]
for f in contratos_features:
    add_icon_text('', f, size=10, color=DARKGRAY)

doc.add_page_break()

# ================================================================
# PAGE 7 - ORGANIGRAMA
# ================================================================
add_centered('Organigrama Interactivo', size=18, color=DARK, bold=True, space_after=8)
add_img('02_personal/06_organigrama.png', width=Inches(6.5), caption='Estructura jerarquica con colores por nivel')
doc.add_paragraph()

add_centered('Portal del Trabajador — Autoservicio', size=18, color=DARK, bold=True, space_after=8)
add_centered('Cada colaborador accede a su informacion desde cualquier dispositivo', size=11, color=GRAY, space_after=12)
add_img('15_portal/01_portal_main.png', width=Inches(6.5), caption='Mi Portal — boletas, vacaciones, asistencia, papeletas')

doc.add_page_break()

# ================================================================
# PAGE 8 - INTELIGENCIA ARTIFICIAL
# ================================================================
add_centered('Inteligencia Artificial Integrada', size=22, color=DARK, bold=True, space_after=12)
add_centered('La unica plataforma de RRHH en Peru con IA nativa', size=14, color=TEAL, bold=True, space_after=16)

ia_features = [
    ('Asistente Virtual', 'Responde consultas de personal, normativa, asistencia. Disponible 24/7 desde cualquier pantalla.'),
    ('Analisis de Contratos', 'Revisa clausulas, identifica riesgos legales, sugiere mejoras segun DL 728.'),
    ('Generacion de Documentos', 'Crea contratos, adendas, prorrogas automaticamente con datos del sistema.'),
    ('Base de Conocimiento', '59 articulos de normativa peruana indexados con IA (RAG + pgvector).'),
    ('Reglas Inteligentes', 'Configura reglas especiales de asistencia conversando con la IA.'),
    ('Dashboard Predictivo', 'Analisis de rotacion, riesgo de fuga, metricas de clima laboral.'),
]

for title, desc in ia_features:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = TEAL
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Cm(0.3)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

doc.add_page_break()

# ================================================================
# PAGE 9 - INTEGRACIONES PERU
# ================================================================
add_centered('Integraciones con entidades peruanas', size=18, color=DARK, bold=True, space_after=8)
add_centered('Exporta en un click, sin errores, en formato oficial', size=12, color=GRAY, space_after=16)

integ = [
    ('SUNAT', 'PLAME mensual, T-Registro altas/bajas, PDT 601'),
    ('AFP', 'AFP Net: Habitat, Integra, Prima, Profuturo'),
    ('ONP', 'Aportes mensuales formato SUNAT'),
    ('Bancos', 'Telecredito BCP, BBVA Continental, Interbank, Scotiabank'),
    ('CTS', 'Depositos semestrales en formato de cada banco'),
    ('EsSalud', 'Aportes patronales mensuales'),
    ('Contabilidad', 'CONCAR, SIGO, SAP, SIRE'),
    ('SCTR', 'Seguro Complementario de Trabajo de Riesgo'),
]

table = doc.add_table(rows=len(integ), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (entity, desc) in enumerate(integ):
    cell0 = table.cell(i, 0)
    cell1 = table.cell(i, 1)
    shade_cell(cell0, '0F766E')
    cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = cell0.paragraphs[0].add_run(entity)
    run0.font.size = Pt(10)
    run0.font.bold = True
    run0.font.color.rgb = WHITE
    run1 = cell1.paragraphs[0].add_run(desc)
    run1.font.size = Pt(10)
    run1.font.color.rgb = DARKGRAY

doc.add_paragraph()
add_img('13_integraciones/01_panel.png', width=Inches(6), caption='Centro de Integraciones y Exportaciones')

doc.add_page_break()

# ================================================================
# PAGE 10 - PRECIOS
# ================================================================
add_centered('Planes y Precios', size=22, color=DARK, bold=True, space_after=16)

planes = [
    ('ASISTENCIA', 'S/ 99', '/mes', [
        'Hasta 50 colaboradores',
        'Tareo, papeletas, horas extra',
        'Banco de horas y reportes',
        'Vacaciones (gestion admin)',
        'Soporte por email',
    ]),
    ('PLANILLA', 'S/ 199', '/mes', [
        'Hasta 100 colaboradores',
        'Todo Asistencia +',
        'Nominas, boletas y liquidacion',
        'SUNAT: PLAME / AFPNet / T-Registro',
        'Contratos y prestamos',
    ]),
    ('TALENTO', 'S/ 399', '/mes', [
        'Hasta 300 colaboradores',
        'Todo Planilla +',
        'Evaluaciones y capacitaciones',
        'Encuestas, clima y organigrama',
        'Analytics y Dashboard Ejecutivo',
    ]),
    ('SUITE', 'S/ 699', '/mes', [
        '300+ colaboradores',
        'Todo Talento +',
        'Reclutamiento y banco de talento',
        'Portal del empleado e IA',
        'Multi-empresa y soporte prioritario',
    ]),
]

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (name, price, period, features) in enumerate(planes):
    cell = table.cell(0, i)
    is_featured = i == 1
    bg = '0F766E' if is_featured else 'F9FAFB'
    shade_cell(cell, bg)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n{name}\n\n')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = WHITE if is_featured else TEAL

    run2 = p.add_run(f'{price}')
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.color.rgb = WHITE if is_featured else DARK

    if period:
        run3 = p.add_run(f'{period}\n\n')
        run3.font.size = Pt(10)
        run3.font.color.rgb = RGBColor(0xC8,0xFA,0xF2) if is_featured else GRAY
    else:
        run3 = p.add_run('\n\n')
        run3.font.size = Pt(10)

    for feat in features:
        run_f = p.add_run(f'  {feat}\n')
        run_f.font.size = Pt(9)
        run_f.font.color.rgb = RGBColor(0xE8,0xFA,0xF7) if is_featured else DARKGRAY

    p.add_run('\n')

doc.add_paragraph()
add_centered('Todos los planes incluyen actualizaciones gratuitas y backup diario', size=10, color=GRAY)

doc.add_page_break()

# ================================================================
# PAGE 11 - POR QUE HARMONI
# ================================================================
add_centered('Por que elegir Harmoni', size=22, color=DARK, bold=True, space_after=16)

reasons = [
    ('100% Peruano', 'Disenado por peruanos para empresas peruanas. Toda la normativa laboral vigente: DL 728, DL 713, Ley 27735, DS 003-97-TR.'),
    ('22 Modulos en Uno', 'No necesitas 5 sistemas diferentes. Personal, asistencia, nominas, vacaciones, reclutamiento, evaluaciones — todo integrado.'),
    ('IA que Trabaja por Ti', 'Asistente virtual, analisis de contratos, reglas inteligentes, dashboard predictivo. La IA mas avanzada en RRHH del Peru.'),
    ('En la Nube o tu Servidor', 'SaaS listo para usar o instalacion on-premise. Tu data es tuya, siempre.'),
    ('Soporte Real', 'No somos un call center. Somos ingenieros que entienden tu negocio y responden en minutos.'),
    ('Precio Justo', 'Desde S/149/mes. Sin contratos de permanencia. Sin costos ocultos. Sin sorpresas.'),
]

for title, desc in reasons:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = TEAL
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

doc.add_page_break()

# ================================================================
# PAGE 12 - CONTACTO / CTA
# ================================================================
for _ in range(5):
    doc.add_paragraph()

add_centered('Empieza hoy', size=28, color=DARK, bold=True, space_after=8)
add_centered('Solicita tu demo gratuita', size=18, color=TEAL, bold=True, space_after=24)

add_divider()
doc.add_paragraph()

add_centered('www.harmoni.pe', size=18, color=TEAL, bold=True, space_after=8)
add_centered('+51 977 538 028', size=16, color=DARK, bold=True, space_after=4)
add_centered('ventas@harmoni.pe', size=14, color=GRAY, space_after=16)

add_divider()
doc.add_paragraph()

add_centered('Harmoni ERP  |  Consorcio Stiler-Ripcon', size=11, color=GRAY, space_after=4)
add_centered('Lima, Peru  |  2026', size=10, color=GRAY)

# ================================================================
# GUARDAR
# ================================================================
out = r"D:\Harmoni\docs\BROCHURE COMERCIAL - HARMONI ERP.docx"
doc.save(out)
print(f"Brochure DOCX: {out}")

# Convert to PDF
try:
    from docx2pdf import convert
    pdf_out = out.replace('.docx', '.pdf')
    convert(out, pdf_out)
    print(f"Brochure PDF: {pdf_out}")
except Exception as e:
    print(f"PDF conversion: {e}")

print("BROCHURE COMPLETADO")
