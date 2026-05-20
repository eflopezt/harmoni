"""
Actualiza el manual agregando secciones faltantes y regenerando PDF.
Lee el DOCX existente y agrega contenido antes del final.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\Harmoni\docs\manual_capturas"
doc_path = r"D:\Harmoni\docs\MANUAL DE USUARIO - HARMONI ERP.docx"

doc = Document(doc_path)

TEAL = RGBColor(0x0F, 0x76, 0x6E)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x6B, 0x72, 0x80)

def add_img(path, width=Inches(6.5), caption=None):
    full = os.path.join(OUT, path) if not os.path.isabs(path) else path
    if os.path.exists(full):
        doc.add_picture(full, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = GRAY

# ================================================================
# AGREGAR SECCIONES FALTANTES AL FINAL
# ================================================================

# --- Firma Digital ---
doc.add_page_break()
doc.add_heading('Apendice A — Firma Digital de Documentos', level=1)
doc.add_paragraph(
    'Harmoni integra dos sistemas de firma electronica: ZapSign (firma digital certificada '
    'con validez legal) y firma interna (firma manuscrita digitalizada en tablet/movil). '
    'Ambos se integran con el modulo de Documentos para firmar contratos, memorandums, '
    'constancias y cualquier documento del legajo.'
)
add_img('09_documentos/04_firma_digital.png', caption='Figura A.1 - Panel de Firma Digital (ZapSign)')
add_img('09_documentos/05_firma_interna.png', caption='Figura A.2 - Firma Interna (Signature Pad)')

# --- Boletas de pago (admin) ---
doc.add_heading('Apendice B — Gestion de Boletas de Pago', level=1)
doc.add_paragraph(
    'El administrador puede subir, publicar y distribuir boletas de pago desde el modulo '
    'Documentos > Boletas. Las boletas se publican y quedan disponibles en el Portal del '
    'Trabajador para descarga individual. Incluye confirmacion de lectura obligatoria.'
)
add_img('09_documentos/03_boletas.png', caption='Figura B.1 - Panel de administracion de boletas')

# --- Nominas avanzadas ---
doc.add_heading('Apendice C — Funciones Avanzadas de Nominas', level=1)

doc.add_heading('C.1 Flujo de Caja', level=2)
doc.add_paragraph(
    'El modulo de Flujo de Caja permite presupuestar y controlar el gasto de planilla '
    'mensual. Compara presupuesto vs ejecutado, proyecta los proximos meses, y alerta '
    'cuando el gasto supera el presupuesto asignado.'
)
add_img('04_nominas/06_flujo_caja.png', caption='Figura C.1 - Flujo de Caja')

doc.add_heading('C.2 Liquidaciones por Cese', level=2)
doc.add_paragraph(
    'Cuando un trabajador es cesado, Harmoni calcula automaticamente su liquidacion: '
    'CTS trunca, gratificacion trunca, vacaciones truncas, compensacion por tiempo '
    'de servicios, y cualquier saldo pendiente. Genera la hoja de liquidacion en PDF '
    'con todos los calculos detallados.'
)
add_img('04_nominas/07_liquidaciones.png', caption='Figura C.2 - Panel de Liquidaciones')

doc.add_heading('C.3 IR 5ta Categoria', level=2)
doc.add_paragraph(
    'Panel dedicado al Impuesto a la Renta de 5ta Categoria: proyeccion anual por trabajador, '
    'calculo mensual segun escala progresiva (8%, 14%, 17%, 20%, 30%), ajustes por gratificacion '
    'y bonificacion extraordinaria. Cumple con la normativa SUNAT vigente.'
)
add_img('04_nominas/08_ir5ta.png', caption='Figura C.3 - Panel IR 5ta Categoria')

# --- Evaluaciones avanzadas ---
doc.add_heading('Apendice D — Evaluaciones y OKRs', level=1)
doc.add_paragraph(
    'Harmoni incluye un sistema completo de evaluacion del desempeno con multiples metodologias: '
    'evaluacion 360 grados, OKR (Objectives & Key Results), matriz 9-Box para talent mapping, '
    'competencias por cargo, y planes de desarrollo individual.'
)
add_img('18_evaluaciones/02_dashboard.png', caption='Figura D.1 - Dashboard de Evaluaciones')
add_img('18_evaluaciones/03_okrs.png', caption='Figura D.2 - Objetivos y Key Results (OKR)')

# --- Reclutamiento Pipeline ---
doc.add_heading('Apendice E — Pipeline de Reclutamiento', level=1)
doc.add_paragraph(
    'El modulo de Reclutamiento incluye un pipeline visual tipo Kanban donde los candidatos '
    'avanzan por etapas configurables: Postulacion, Filtro CV, Entrevista RRHH, Entrevista '
    'Tecnica, Oferta, Contratacion. Cada etapa tiene scoring automatico y notas del evaluador.'
)
add_img('06_reclutamiento/03_pipeline.png', caption='Figura E.1 - Pipeline Kanban de Candidatos')

# --- Disciplinaria Dashboard ---
doc.add_heading('Apendice F — Dashboard Disciplinario', level=1)
add_img('19_disciplinaria/02_dashboard.png', caption='Figura F.1 - Dashboard de Gestion Disciplinaria')

# --- Viaticos Dashboard ---
doc.add_heading('Apendice G — Dashboard de Viaticos', level=1)
add_img('21_viaticos/02_dashboard.png', caption='Figura G.1 - Dashboard de Viaticos y Gastos')

# --- Salarios ---
doc.add_heading('Apendice H — Estructura Salarial y Equidad', level=1)
doc.add_paragraph(
    'Harmoni permite definir bandas salariales por cargo y nivel, analizar la equidad '
    'interna (brecha salarial por genero, area, antiguedad), simular incrementos masivos '
    'y comparar con el mercado.'
)
add_img('22_salarios/02_bandas.png', caption='Figura H.1 - Bandas Salariales')
add_img('22_salarios/03_equidad.png', caption='Figura H.2 - Analisis de Equidad Salarial')

# --- Administracion del Sistema ---
doc.add_heading('Apendice I — Administracion del Sistema', level=1)

doc.add_heading('I.1 Gestion de Usuarios', level=2)
doc.add_paragraph(
    'Administracion completa de usuarios: crear, editar, vincular a empleado, asignar '
    'perfiles de acceso. Incluye impersonacion (ver el sistema como otro usuario) para '
    'soporte tecnico. Acciones masivas para asignacion de permisos.'
)
add_img('24_sistema/02_gestion_usuarios.png', caption='Figura I.1 - Gestion de Usuarios')

doc.add_heading('I.2 Control de Accesos (RBAC)', level=2)
doc.add_paragraph(
    'Sistema de control de acceso basado en roles: perfiles predefinidos (Admin, RRHH, '
    'Jefatura, Empleado) con permisos granulares por modulo. Cada usuario ve solo lo que '
    'le corresponde segun su perfil.'
)
add_img('24_sistema/03_accesos.png', caption='Figura I.2 - Control de Accesos RBAC')

doc.add_heading('I.3 Auditoria de Cambios', level=2)
doc.add_paragraph(
    'Registro automatico de todos los cambios realizados en el sistema: quien cambio que, '
    'cuando, valor anterior y nuevo. Timeline por registro. Exportable para cumplimiento '
    'regulatorio. Busqueda avanzada por usuario, fecha, modelo y accion.'
)
add_img('24_sistema/04_auditoria.png', caption='Figura I.3 - Log de Auditoria')

# --- Vacaciones calendario ---
doc.add_heading('Apendice J — Calendario de Vacaciones', level=1)
add_img('05_vacaciones/04_calendario.png', caption='Figura J.1 - Calendario de Vacaciones del Equipo')

# --- Plantillas de Contrato ---
doc.add_heading('Apendice K — Plantillas de Contrato', level=1)
doc.add_paragraph(
    'Harmoni incluye 11 plantillas de contrato basadas en la normativa laboral peruana: '
    'Contrato por Obra (Confianza, Fiscalizable, Generico), Contrato Indefinido, '
    'Contrato por Necesidad de Mercado, Prorroga, y 3 tipos de Adenda (Aumento Salarial, '
    'Cambio de Cargo, Cambio de Condiciones). Cada plantilla usa 30+ placeholders que se '
    'resuelven automaticamente con datos del sistema.'
)
add_img('02_personal/13_plantillas_contrato.png', caption='Figura K.1 - Gestion de Plantillas de Contrato')

# --- Analytics adicional ---
doc.add_heading('Apendice L — Analytics Avanzado', level=1)
add_img('11_analytics/03_headcount.png', caption='Figura L.1 - Analisis de Headcount')

# ================================================================
# GUARDAR
# ================================================================
doc.save(doc_path)
print(f"Manual actualizado: {doc_path}")

# PDF
from docx2pdf import convert
pdf_path = doc_path.replace('.docx', '.pdf')
convert(doc_path, pdf_path)
print(f"PDF actualizado: {pdf_path}")

# Stats
import subprocess
result = subprocess.run(['ls', '-lh', doc_path], capture_output=True, text=True)
print(result.stdout.strip())
result2 = subprocess.run(['ls', '-lh', pdf_path], capture_output=True, text=True)
print(result2.stdout.strip())
print("MANUAL ACTUALIZADO EXITOSAMENTE")
